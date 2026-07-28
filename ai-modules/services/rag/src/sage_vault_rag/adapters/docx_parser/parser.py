import io
import logging
import re
from zipfile import BadZipFile

from docx import Document
from docx.document import Document as _DocumentObject
from docx.opc.exceptions import PackageNotFoundError
from docx.text.paragraph import Paragraph

from sage_vault_rag.model.parsed_document import ParsedDocument, ParsedParagraph

logger = logging.getLogger(__name__)

_HEADING_STYLE = re.compile(r"^Heading\s+([1-6])$")


class DocxParser:
    """DOCX 解析器，保留段落与标题层级。

    失败语义（与 03d 工单对齐）：
    - 损坏或无法识别的 DOCX（非 ZIP、非 OOXML 包）抛 ``ValueError("文件损坏，无法解析: <filename>")``。
    - 空内容或有效但无可提取文本的 DOCX 抛 ``ValueError("文档内容为空: <filename>")``。

    成功解析时按文档顺序遍历 ``document.paragraphs``：
    - 样式为 ``Heading 1`` ~ ``Heading 6`` 的段落作为标题段输出，``text`` 与 ``heading``
      均为标题文本，并成为后续正文段的当前标题。Heading 7~9 在 Word 默认 UI 中不暴露，
      且与 Markdown ATX 六级标题对齐，故不识别。
    - 非标题段落输出 ``text`` 为正文，``heading`` 为最近一个标题文本（无标题时为 ``None``）。
    - DOCX 无页码概念，``page_number`` 始终为 ``None``。
    空白段落（无文本）被跳过，不产生 ``ParsedParagraph``。
    """

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        if not content:
            raise ValueError(f"文档内容为空: {filename}")
        document = self._open_document(content, filename)
        paragraphs = self._extract_paragraphs(document, filename)
        if not paragraphs:
            raise ValueError(f"文档内容为空: {filename}")
        return ParsedDocument(paragraphs=paragraphs)

    @staticmethod
    def _open_document(content: bytes, filename: str) -> _DocumentObject:
        try:
            return Document(io.BytesIO(content))
        except (BadZipFile, PackageNotFoundError) as exception:
            raise ValueError(f"文件损坏，无法解析: {filename}") from exception

    @staticmethod
    def _extract_paragraphs(document: _DocumentObject, filename: str) -> list[ParsedParagraph]:
        paragraphs: list[ParsedParagraph] = []
        current_heading: str | None = None
        try:
            doc_paragraphs: list[Paragraph] = list(document.paragraphs)
        except (KeyError, IndexError) as exception:
            logger.warning("DOCX 段落读取失败: filename=%s", filename, exc_info=True)
            raise ValueError(f"文件损坏，无法解析: {filename}") from exception
        for paragraph in doc_paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if _HEADING_STYLE.match(style_name):
                current_heading = text
            paragraphs.append(ParsedParagraph(text=text, heading=current_heading))
        return paragraphs
