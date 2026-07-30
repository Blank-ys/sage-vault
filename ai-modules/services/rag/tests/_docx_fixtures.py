"""DOCX 测试夹具生成工具。

使用 python-docx 生成测试用 DOCX 字节内容，避免引入额外测试依赖。
所有夹具均为内存中构造，返回字节串，便于直接喂给 ``DocxParser.parse``。
"""

import io
from collections.abc import Sequence

from docx import Document
from docx.document import Document as _DocumentObject


def _save_to_bytes(document: _DocumentObject) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_text_docx(text: str) -> bytes:
    """生成包含单个正文段的 DOCX。"""
    document = Document()
    document.add_paragraph(text)
    return _save_to_bytes(document)


def make_multi_paragraph_docx(paragraphs: Sequence[str]) -> bytes:
    """生成包含多个正文段的 DOCX，按顺序写入。"""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    return _save_to_bytes(document)


def make_docx_with_headings(sections: Sequence[tuple[int, str, Sequence[str]]]) -> bytes:
    """生成带标题层级的 DOCX。

    ``sections`` 每项为 ``(heading_level, heading_text, body_paragraphs)``：
    - ``heading_level`` 取 1~6，对应 Heading 1 ~ Heading 6 样式；
    - ``heading_text`` 为标题正文；
    - ``body_paragraphs`` 为该标题下的正文段列表。
    """
    document = Document()
    for level, heading_text, body_paragraphs in sections:
        document.add_heading(heading_text, level=level)
        for body in body_paragraphs:
            document.add_paragraph(body)
    return _save_to_bytes(document)


def make_empty_docx() -> bytes:
    """生成有效但无任何段落文本的 DOCX。"""
    document = Document()
    return _save_to_bytes(document)


def make_whitespace_only_docx() -> bytes:
    """生成仅含空白段落文本的 DOCX。"""
    document = Document()
    document.add_paragraph("   ")
    document.add_paragraph("\n\n")
    return _save_to_bytes(document)


def make_docx_with_leading_body(leading_body: str, sections: Sequence[tuple[int, str, Sequence[str]]]) -> bytes:
    """生成在首个标题之前存在正文段的 DOCX，用于验证无标题正文的 heading 为 None。

    ``leading_body`` 为首个标题前的正文段；``sections`` 与 ``make_docx_with_headings`` 一致。
    """
    document = Document()
    document.add_paragraph(leading_body)
    for level, heading_text, body_paragraphs in sections:
        document.add_heading(heading_text, level=level)
        for body in body_paragraphs:
            document.add_paragraph(body)
    return _save_to_bytes(document)
